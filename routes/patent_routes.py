from flask import Blueprint, render_template, request, session, redirect, url_for
from flask import Response, send_file
import os

from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer

from services.lens_service import search_lens_patents
from services.prior_art_service import analyze_prior_art
from services.gemini_service import generate_patent_draft
from services.diagram_service import save_uploaded_diagram, generate_drawing_text_from_uploaded_image
from utils.prompt_builder import build_patent_prompt
from models.draft_model import DraftModel
from models.related_patent_model import RelatedPatentModel


patent = Blueprint("patent", __name__)

draft_model = DraftModel()
related_patent_model = RelatedPatentModel()


def ensure_user_draft_folder(user_id):
    folder = os.path.join("saved_drafts", f"user_{user_id}")
    os.makedirs(folder, exist_ok=True)
    return folder


@patent.route("/questionnaire")
def questionnaire():
    if "user_id" not in session:
        return redirect("/login")
    return render_template("questionnaire.html")


@patent.route("/generate_draft", methods=["POST"])
def generate_draft():
    if "user_id" not in session:
        return redirect("/login")

    try:
        data = {
            "title": request.form.get("title", "").strip(),
            "field": request.form.get("field", "").strip(),
            "problem": request.form.get("problem", "").strip(),
            "existing_solution": request.form.get("existing_solution", "").strip(),
            "proposed_solution": request.form.get("proposed_solution", "").strip(),
            "working": request.form.get("working", "").strip(),
            "advantages": request.form.get("advantages", "").strip(),
            "components": request.form.get("components", "").strip(),
            "applications": request.form.get("applications", "").strip()
        }

        uploaded_file = request.files.get("uploaded_diagram")
        diagram_path = save_uploaded_diagram(uploaded_file)

        diagram_caption = ""
        drawing_description = ""
        if diagram_path:
            diagram_caption, drawing_description = generate_drawing_text_from_uploaded_image(
                data,
                diagram_path
            )

        lens_response = search_lens_patents(data)

        raw_related_patents = lens_response.get("results", [])
        lens_error = lens_response.get("error", "")
        lens_payload = lens_response.get("payload", {})
        lens_query = lens_response.get("query_text", "")
        lens_queries = lens_response.get("attempted_queries", [])
        lens_debug_info = lens_response.get("debug_info", [])
        lens_fingerprint = lens_response.get("fingerprint", {})

        prior_art_analysis = analyze_prior_art(data, raw_related_patents)
        related_patents = prior_art_analysis.get("ranked_patents", raw_related_patents)

        related_patents_text = ""
        for i, item in enumerate(related_patents, start=1):
            related_patents_text += (
                f"{i}. Title: {item.get('patent_title', '')}\n"
                f"   Number: {item.get('patent_number', '')}\n"
                f"   URL: {item.get('patent_url', '')}\n"
                f"   Similarity: {item.get('similarity_percent', 0)}%\n"
                f"   Abstract: {item.get('abstract_text', '')}\n\n"
            )

        prompt_text = build_patent_prompt(
            data=data,
            related_patents_text=related_patents_text,
            drawing_description=drawing_description,
            prior_art_analysis=prior_art_analysis.get("formatted_text", "")
        )

        draft_text = generate_patent_draft(prompt_text)

        draft_id = draft_model.create_draft(
            user_id=session["user_id"],
            title=data["title"],
            field_of_technology=data["field"],
            draft_content=draft_text,
            prior_art_analysis=prior_art_analysis.get("formatted_text", ""),
            diagram_path=diagram_path,
            diagram_caption=diagram_caption,
            drawing_description=drawing_description
        )

        # Save first version
        draft_model.save_draft_version(draft_id, 1, draft_text)

        if related_patents:
            related_patent_model.save_related_patents(draft_id, related_patents)

        diagram_url = ""
        if diagram_path and diagram_path.startswith("uploads/"):
            relative_filename = diagram_path[len("uploads/"):]
            diagram_url = url_for("uploaded_file", filename=relative_filename)

        return render_template(
            "generated_draft.html",
            draft=draft_text,
            draft_id=draft_id,
            related_patents=related_patents,
            prior_art_analysis=prior_art_analysis,
            diagram_path=diagram_path,
            diagram_url=diagram_url,
            diagram_caption=diagram_caption,
            drawing_description=drawing_description,
            lens_query=lens_query,
            lens_queries=lens_queries,
            lens_error=lens_error,
            lens_payload=lens_payload,
            lens_debug_info=lens_debug_info,
            lens_fingerprint=lens_fingerprint
        )

    except Exception as e:
        return render_template(
            "generated_draft.html",
            draft="Patent draft could not be generated due to an internal error.",
            draft_id=0,
            related_patents=[],
            prior_art_analysis={
                "status": "System Error",
                "similarity_score": 0,
                "analysis_text": f"An error occurred while processing the request: {str(e)}",
                "formatted_text": f"STATUS: System Error\nSIMILARITY_SCORE: 0%\nEXPLANATION: {str(e)}"
            },
            diagram_path="",
            diagram_url="",
            diagram_caption="",
            drawing_description="",
            lens_query="",
            lens_queries=[],
            lens_error=str(e),
            lens_payload={},
            lens_debug_info=[],
            lens_fingerprint={}
        )


@patent.route("/edit_draft/<int:draft_id>", methods=["GET", "POST"])
def edit_draft(draft_id):
    if "user_id" not in session:
        if request.method == "POST":
            return {"success": False, "message": "Unauthorized"}, 401
        return redirect("/login")

    draft_record = draft_model.get_draft_by_id(draft_id)
    if not draft_record:
        if request.method == "POST":
            return {"success": False, "message": "Draft not found"}, 404
        return "Draft not found", 404
        
    if draft_record["user_id"] != session["user_id"]:
        if request.method == "POST":
            return {"success": False, "message": "Forbidden: You do not own this draft"}, 403
        return "Forbidden: You do not own this draft", 403

    if request.method == "POST":
        new_text = request.form.get("draft_text", "").strip()
        if not new_text:
            return {"success": False, "message": "Draft text cannot be empty"}, 400

        draft_model.update_draft_text(draft_id, new_text, True)
        return {"success": True, "message": "Draft updated successfully"}

    return render_template("edit_draft.html", draft=draft_record)


@patent.route("/regenerate_draft/<int:draft_id>", methods=["POST"])
def regenerate_draft(draft_id):
    if "user_id" not in session:
        return {"success": False, "message": "Unauthorized"}, 401

    draft_record = draft_model.get_draft_by_id(draft_id)
    if not draft_record:
        return {"success": False, "message": "Draft not found"}, 404
        
    if draft_record["user_id"] != session["user_id"]:
        return {"success": False, "message": "Forbidden: You do not own this draft"}, 403

    # Rebuild prompt from original stored data
    data = {
        "title": draft_record.get("title", ""),
        "field": draft_record.get("field", ""),
        "problem": draft_record.get("problem", ""),
        "existing_solution": draft_record.get("existing_solution", ""),
        "proposed_solution": draft_record.get("proposed_solution", ""),
        "working": draft_record.get("working", ""),
        "advantages": draft_record.get("advantages", ""),
        "components": draft_record.get("components", ""),
        "applications": draft_record.get("applications", "")
    }

    related_patents_text = ""
    drawing_description = draft_record.get("drawing_description", "")
    prior_art_analysis = draft_record.get("prior_art_analysis", "")

    prompt_text = build_patent_prompt(
        data=data,
        related_patents_text=related_patents_text,
        drawing_description=drawing_description,
        prior_art_analysis=prior_art_analysis
    )

    new_draft = generate_patent_draft(prompt_text)

    next_version = draft_model.get_next_version_no(draft_id)
    draft_model.save_draft_version(draft_id, next_version, new_draft)
    draft_model.update_draft_text(draft_id, new_draft, False)
    draft_model.update_draft_version_no(draft_id, next_version)

    return {
        "success": True,
        "message": f"Draft regenerated as version {next_version}",
        "draft_text": new_draft
    }


@patent.route("/download_draft/txt/<int:draft_id>")
def download_draft_txt(draft_id):
    if "user_id" not in session:
        return redirect("/login")

    draft_record = draft_model.get_draft_by_id(draft_id)
    if not draft_record:
        return "Draft not found", 404
        
    if draft_record["user_id"] != session["user_id"]:
        return "Forbidden: You do not own this draft", 403

    return Response(
        draft_record["draft_text"],
        mimetype="text/plain",
        headers={"Content-Disposition": f"attachment; filename=draft_{draft_id}.txt"}
    )


@patent.route("/download_draft/docx/<int:draft_id>")
def download_draft_docx(draft_id):
    if "user_id" not in session:
        return redirect("/login")

    draft_record = draft_model.get_draft_by_id(draft_id)
    if not draft_record:
        return "Draft not found", 404
        
    if draft_record["user_id"] != session["user_id"]:
        return "Forbidden: You do not own this draft", 403

    folder = ensure_user_draft_folder(session["user_id"])
    filepath = os.path.join(folder, f"draft_{draft_id}.docx")

    doc = Document()
    doc.add_heading(draft_record["title"], 0)
    doc.add_paragraph(draft_record["draft_text"])
    doc.save(filepath)

    return send_file(filepath, as_attachment=True)


@patent.route("/download_draft/pdf/<int:draft_id>")
def download_draft_pdf(draft_id):
    if "user_id" not in session:
        return redirect("/login")

    draft_record = draft_model.get_draft_by_id(draft_id)
    if not draft_record:
        return "Draft not found", 404
        
    if draft_record["user_id"] != session["user_id"]:
        return "Forbidden: You do not own this draft", 403

    folder = ensure_user_draft_folder(session["user_id"])
    filepath = os.path.join(folder, f"draft_{draft_id}.pdf")

    # Use SimpleDocTemplate for professional rendering with word wrapping
    doc = SimpleDocTemplate(filepath, pagesize=A4, rightMargin=50, leftMargin=50, topMargin=50, bottomMargin=50)
    styles = getSampleStyleSheet()
    
    title_style = ParagraphStyle(
        'DocTitle',
        parent=styles['Heading1'],
        fontName='Helvetica-Bold',
        fontSize=18,
        leading=22,
        spaceAfter=15
    )
    body_style = ParagraphStyle(
        'DocBody',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=10,
        leading=14,
        spaceAfter=8
    )

    story = []
    story.append(Paragraph(draft_record["title"] or "Untitled Patent Draft", title_style))
    story.append(Spacer(1, 10))

    # Split lines by newlines and add paragraph blocks
    paragraphs = (draft_record["draft_text"] or "").split("\n\n")
    for p_text in paragraphs:
        p_text_html = p_text.strip().replace("\n", "<br/>")
        if p_text_html:
            story.append(Paragraph(p_text_html, body_style))
            story.append(Spacer(1, 6))

    doc.build(story)
    return send_file(filepath, as_attachment=True)


@patent.route("/draft_versions/<int:draft_id>")
def draft_versions(draft_id):
    if "user_id" not in session:
        return redirect("/login")

    draft_record = draft_model.get_draft_by_id(draft_id)
    if not draft_record:
        return "Draft not found", 404
        
    if draft_record["user_id"] != session["user_id"]:
        return "Forbidden: You do not own this draft", 403

    versions = draft_model.get_draft_versions(draft_id)
    return render_template("draft_versions.html", versions=versions, draft_id=draft_id)